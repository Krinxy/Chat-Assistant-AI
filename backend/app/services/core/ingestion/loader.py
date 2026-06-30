from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

_TXT_SUFFIXES = {".txt", ".md"}
_PDF_SUFFIX = ".pdf"
_DOCX_SUFFIX = ".docx"
_SUPPORTED_SUFFIXES = _TXT_SUFFIXES | {_PDF_SUFFIX, _DOCX_SUFFIX}


class UnsupportedDocumentError(ValueError):
    """Raised when a document has a file type the loader cannot extract text from."""


@dataclass(frozen=True)
class LoadedDocument:
    source: str
    text: str


class DocumentLoader:
    @staticmethod
    def supported_suffixes() -> frozenset[str]:
        return frozenset(_SUPPORTED_SUFFIXES)

    @staticmethod
    def extract_text(filename: str, data: bytes) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix in _TXT_SUFFIXES:
            return DocumentLoader._extract_txt(data)
        if suffix == _PDF_SUFFIX:
            return DocumentLoader._extract_pdf(data)
        if suffix == _DOCX_SUFFIX:
            return DocumentLoader._extract_docx(data)
        raise UnsupportedDocumentError(f"Unsupported document type: '{suffix or filename}'")

    @staticmethod
    def extract_segments(filename: str, data: bytes) -> list[str]:
        """Extract a document into logical segments for structure-aware chunking.

        Unlike :meth:`extract_text` (which flattens everything into one string), this keeps
        table rows as separate, self-contained segments — each row becomes ``"Header: value | …"``
        so a downstream chunker can store one table entry per chunk. PDFs are parsed with
        ``pdfplumber`` for true table detection; plain text and DOCX yield a single segment.
        """
        suffix = Path(filename).suffix.lower()
        if suffix in _TXT_SUFFIXES:
            return DocumentLoader._non_empty([DocumentLoader._extract_txt(data)])
        if suffix == _PDF_SUFFIX:
            return DocumentLoader._extract_pdf_segments(data)
        if suffix == _DOCX_SUFFIX:
            return DocumentLoader._non_empty([DocumentLoader._extract_docx(data)])
        raise UnsupportedDocumentError(f"Unsupported document type: '{suffix or filename}'")

    @staticmethod
    def load_document(path: str | Path) -> LoadedDocument:
        resolved = Path(path)
        data = resolved.read_bytes()
        text = DocumentLoader.extract_text(resolved.name, data)
        return LoadedDocument(source=resolved.name, text=text)

    @staticmethod
    def _extract_txt(data: bytes) -> str:
        return data.decode("utf-8", errors="replace").strip()

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(page.strip() for page in pages if page.strip()).strip()

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        document = Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(paragraph.strip() for paragraph in paragraphs if paragraph.strip()).strip()

    @staticmethod
    def _extract_pdf_segments(data: bytes) -> list[str]:
        """Return one segment per table row plus the page's non-table prose.

        Falls back to flat :meth:`_extract_pdf` (pypdf) when pdfplumber is unavailable or finds
        no tables, so non-tabular PDFs keep working exactly as before.
        """
        try:
            import pdfplumber
        except ImportError:
            return DocumentLoader._non_empty([DocumentLoader._extract_pdf(data)])

        segments: list[str] = []
        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                tables = page.find_tables()
                for table in tables:
                    segments.extend(DocumentLoader._serialize_table(table.extract()))
                prose = DocumentLoader._page_prose_outside_tables(page, tables)
                if prose:
                    segments.append(prose)

        segments = DocumentLoader._non_empty(segments)
        if segments:
            return segments
        # No table or word geometry recovered — fall back to flat extraction.
        return DocumentLoader._non_empty([DocumentLoader._extract_pdf(data)])

    @staticmethod
    def _serialize_table(rows: list[list[str | None]]) -> list[str]:
        """Turn a table into one ``"Header: value | …"`` string per data row.

        The first non-empty row is treated as the header when every cell is filled; otherwise
        cells are joined positionally. Empty cells are dropped so each row stays self-contained.
        """
        cleaned = [
            [(cell or "").strip().replace("\n", " ") for cell in row]
            for row in rows
            if any((cell or "").strip() for cell in row)
        ]
        if not cleaned:
            return []

        header = cleaned[0]
        has_header = len(cleaned) > 1 and all(cell for cell in header)
        body = cleaned[1:] if has_header else cleaned

        serialized: list[str] = []
        for row in body:
            if has_header:
                parts = [f"{header[i]}: {row[i]}" for i in range(min(len(header), len(row))) if row[i]]
            else:
                parts = [cell for cell in row if cell]
            line = " | ".join(parts)
            if line:
                serialized.append(line)
        return serialized

    @staticmethod
    def _page_prose_outside_tables(page: Any, tables: list[Any]) -> str:
        """Join the page's words that fall outside any detected table bounding box."""
        bboxes = [table.bbox for table in tables]

        def _inside_any(word: dict[str, Any]) -> bool:
            cx = (word["x0"] + word["x1"]) / 2
            cy = (word["top"] + word["bottom"]) / 2
            return any(x0 <= cx <= x1 and top <= cy <= bottom for (x0, top, x1, bottom) in bboxes)

        words = [word["text"] for word in page.extract_words() if not _inside_any(word)]
        return " ".join(words).strip()

    @staticmethod
    def _non_empty(segments: list[str]) -> list[str]:
        return [segment for segment in segments if segment.strip()]


# module-level aliases for backward compat
supported_suffixes = DocumentLoader.supported_suffixes
extract_text = DocumentLoader.extract_text
extract_segments = DocumentLoader.extract_segments
load_document = DocumentLoader.load_document
