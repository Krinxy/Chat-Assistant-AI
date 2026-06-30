from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from backend.app.services.core.ingestion.loader import (
    DocumentLoader,
    LoadedDocument,
    UnsupportedDocumentError,
    extract_segments,
    extract_text,
    load_document,
    supported_suffixes,
)


def test_extract_text_reads_plain_text() -> None:
    assert extract_text("notes.txt", b"  hello world  ") == "hello world"


def test_extract_text_reads_markdown_suffix() -> None:
    assert extract_text("README.md", "# Title\nbody".encode("utf-8")) == "# Title\nbody"


def test_extract_text_reads_docx_paragraphs() -> None:
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("")
    document.add_paragraph("Second paragraph")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text("report.docx", buffer.getvalue())

    assert text == "First paragraph\nSecond paragraph"


def test_extract_text_reads_pdf_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    class _FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class _FakeReader:
        def __init__(self, _stream: object) -> None:
            self.pages = [_FakePage("Page one"), _FakePage("   "), _FakePage("Page two")]

    import backend.app.services.core.ingestion.loader as _loader_mod

    monkeypatch.setattr(_loader_mod, "PdfReader", _FakeReader)

    assert extract_text("manual.pdf", b"%PDF-fake") == "Page one\n\nPage two"


def test_extract_text_rejects_unsupported_suffix() -> None:
    with pytest.raises(UnsupportedDocumentError, match="Unsupported document type"):
        extract_text("archive.zip", b"data")


def test_supported_suffixes_contains_expected_types() -> None:
    suffixes = supported_suffixes()

    assert {".pdf", ".docx", ".txt", ".md"} <= suffixes


def test_serialize_table_pairs_header_with_each_row() -> None:
    rows = [
        ["Modul", "Datum", "Uhrzeit"],
        ["Robotik", "28.07.2026", "11:00"],
        ["Deep Learning", "29.07.2026", "08:00"],
    ]

    serialized = DocumentLoader._serialize_table(rows)

    assert serialized == [
        "Modul: Robotik | Datum: 28.07.2026 | Uhrzeit: 11:00",
        "Modul: Deep Learning | Datum: 29.07.2026 | Uhrzeit: 08:00",
    ]


def test_serialize_table_drops_empty_cells_and_blank_rows() -> None:
    rows = [
        ["Modul", "Datum", "Raum"],
        ["Big Data", "30.07.2026", None],
        [None, None, None],
    ]

    serialized = DocumentLoader._serialize_table(rows)

    assert serialized == ["Modul: Big Data | Datum: 30.07.2026"]


def test_serialize_table_joins_positionally_without_header() -> None:
    # A single data row (no separate header) is joined positionally.
    serialized = DocumentLoader._serialize_table([["Robotik", "28.07.2026"]])

    assert serialized == ["Robotik | 28.07.2026"]


def test_extract_segments_returns_single_segment_for_plain_text() -> None:
    assert extract_segments("notes.txt", b"  hello world  ") == ["hello world"]


def test_extract_segments_rejects_unsupported_suffix() -> None:
    with pytest.raises(UnsupportedDocumentError):
        extract_segments("archive.zip", b"data")


def test_load_document_reads_file_from_disk(tmp_path: Path) -> None:
    source = tmp_path / "doc.txt"
    source.write_text("on-prem document", encoding="utf-8")

    loaded = load_document(source)

    assert loaded == LoadedDocument(source="doc.txt", text="on-prem document")
